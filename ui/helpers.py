from __future__ import annotations

import json
import re
import unicodedata
from pathlib import Path
from typing import Iterable

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


UI_FONT_FAMILY = 'STKaiti, KaiTi, "楷体", serif'
WORD_FONT_NAME = "华文楷体"
WORD_FONT_SIZE = Pt(10.5)


def _safe_filename(filename: str) -> str:
    normalized = unicodedata.normalize("NFKC", Path(filename).name).strip()
    normalized = re.sub(r"[\x00-\x1f\x7f]", "", normalized)
    normalized = re.sub(r"[<>:\"/\\|?*]", "_", normalized)
    normalized = normalized.strip(" .")
    return normalized or "uploaded_document"


def save_uploaded_files(uploaded_files: Iterable, target_dir: str | Path) -> list[Path]:
    base = Path(target_dir)
    base.mkdir(parents=True, exist_ok=True)
    saved_paths: list[Path] = []
    used_names: set[str] = set()
    for file in uploaded_files:
        safe_name = _safe_filename(file.name)
        stem = Path(safe_name).stem
        suffix = Path(safe_name).suffix.lower()
        candidate = safe_name
        sequence = 2
        while candidate.lower() in used_names or (base / candidate).exists():
            candidate = f"{stem}_{sequence}{suffix}"
            sequence += 1
        used_names.add(candidate.lower())
        save_path = base / candidate
        save_path.write_bytes(file.getbuffer())
        saved_paths.append(save_path)
    return saved_paths



def read_json(path: str | Path):
    p = Path(path)
    if not p.exists():
        return None
    return json.loads(p.read_text(encoding="utf-8"))



def read_text(path: str | Path) -> str | None:
    p = Path(path)
    if not p.exists():
        return None
    return p.read_text(encoding="utf-8")



def reset_directory(path: str | Path, suffixes: tuple[str, ...] | None = None) -> None:
    p = Path(path)
    if not p.exists():
        p.mkdir(parents=True, exist_ok=True)
        return
    for child in p.iterdir():
        if child.is_file():
            if suffixes is None or child.suffix.lower() in suffixes:
                child.unlink(missing_ok=True)
        elif child.is_dir() and suffixes is None:
            import shutil
            shutil.rmtree(child, ignore_errors=True)



def pretty_json_block(data) -> None:
    st.code(json.dumps(data, ensure_ascii=False, indent=2), language="json")



def _set_run_font(run, font_name: str = WORD_FONT_NAME, size=WORD_FONT_SIZE, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = font_name
    run.font.size = size
    r = run._element.rPr
    if r is not None:
        r.rFonts.set(qn('w:eastAsia'), font_name)
        r.rFonts.set(qn('w:ascii'), font_name)
        r.rFonts.set(qn('w:hAnsi'), font_name)



def _set_paragraph_font(paragraph, font_name: str = WORD_FONT_NAME, size=WORD_FONT_SIZE) -> None:
    for run in paragraph.runs:
        _set_run_font(run, font_name=font_name, size=size, bold=run.bold)



def _apply_document_style(doc: Document) -> None:
    normal_style = doc.styles['Normal']
    normal_style.font.name = WORD_FONT_NAME
    normal_style.font.size = WORD_FONT_SIZE
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), WORD_FONT_NAME)
    normal_style._element.rPr.rFonts.set(qn('w:ascii'), WORD_FONT_NAME)
    normal_style._element.rPr.rFonts.set(qn('w:hAnsi'), WORD_FONT_NAME)



def _add_md_inline(paragraph, text: str) -> None:
    if not text:
        return
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run)



def export_markdown_to_word(md_path: str | Path, docx_path: str | Path, *, title: str = "结构化尽调结果") -> Path:
    md_file = Path(md_path)
    docx_file = Path(docx_path)
    docx_file.parent.mkdir(parents=True, exist_ok=True)
    text = md_file.read_text(encoding='utf-8')

    doc = Document()
    _apply_document_style(doc)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    _set_run_font(title_run, size=Pt(16), bold=True)

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            _set_run_font(run, size=Pt(15), bold=True)
            continue
        if line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip())
            _set_run_font(run, size=Pt(13), bold=True)
            continue
        if line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            _set_run_font(run, size=Pt(12), bold=True)
            continue
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_md_inline(p, line[2:].strip())
            _set_paragraph_font(p)
            continue
        p = doc.add_paragraph()
        _add_md_inline(p, line)
        _set_paragraph_font(p)

    doc.save(str(docx_file))
    return docx_file
