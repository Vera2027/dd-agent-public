from __future__ import annotations

from pathlib import Path
from typing import Iterator
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document as DocxDocument
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.section import Section
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from dd_agent.domain.schemas import FileInfo, ParsedDocument, TextUnit
from dd_agent.ingest.readers.base import BaseReader

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": _W_NS}


class DocxReader(BaseReader):
    supported_suffixes = (".docx",)

    def read(self, path: Path) -> ParsedDocument:
        file_info = FileInfo.from_path(path)
        try:
            return self._read_with_python_docx(path, file_info)
        except Exception as exc:
            fallback_doc = self._read_with_zip_fallback(path, file_info, original_error=exc)
            return fallback_doc

    def _read_with_python_docx(self, path: Path, file_info: FileInfo) -> ParsedDocument:
        doc = DocxDocument(str(path))

        text_units: list[TextUnit] = []
        paragraph_no = 0
        block_type_counts = {
            "body_paragraph": 0,
            "body_table": 0,
            "header_paragraph": 0,
            "header_table": 0,
            "footer_paragraph": 0,
            "footer_table": 0,
        }

        for block in _iter_block_items(doc):
            extracted = _extract_block_text(block)
            cleaned = _normalize_text(extracted)
            if not cleaned:
                continue
            paragraph_no += 1
            block_kind = _block_kind(block, zone="body")
            block_type_counts[block_kind] += 1
            text_units.append(
                TextUnit(
                    document_name=file_info.document_name,
                    file_path=str(file_info.path),
                    file_type=file_info.file_type,
                    locator_type="paragraph",
                    locator_value=paragraph_no,
                    text=cleaned,
                    metadata={
                        "source_zone": "body",
                        "source_kind": block_kind,
                    },
                )
            )

        for section_index, section in enumerate(doc.sections, start=1):
            paragraph_no = self._append_story_blocks(
                file_info=file_info,
                section=section,
                story_name="header",
                story_container=section.header,
                start_no=paragraph_no,
                text_units=text_units,
                block_type_counts=block_type_counts,
                section_index=section_index,
            )
            paragraph_no = self._append_story_blocks(
                file_info=file_info,
                section=section,
                story_name="footer",
                story_container=section.footer,
                start_no=paragraph_no,
                text_units=text_units,
                block_type_counts=block_type_counts,
                section_index=section_index,
            )

        raw_text = "\n\n".join(unit.text for unit in text_units)
        metadata = self._build_base_metadata(file_info)
        metadata.update(
            {
                "raw_paragraph_count": len(doc.paragraphs),
                "non_empty_paragraph_count": len(text_units),
                "body_table_count": len(doc.tables),
                "section_count": len(doc.sections),
                "block_type_counts": block_type_counts,
                "parse_backend": "python-docx",
            }
        )

        return ParsedDocument(
            document_name=file_info.document_name,
            file_path=str(file_info.path),
            file_type=file_info.file_type,
            raw_text=raw_text,
            text_units=text_units,
            metadata=metadata,
        )

    def _read_with_zip_fallback(self, path: Path, file_info: FileInfo, *, original_error: Exception) -> ParsedDocument:
        text_units: list[TextUnit] = []
        paragraph_no = 0
        block_type_counts = {
            "body_paragraph": 0,
            "body_table": 0,
            "header_paragraph": 0,
            "header_table": 0,
            "footer_paragraph": 0,
            "footer_table": 0,
        }
        raw_paragraph_count = 0
        body_table_count = 0

        with ZipFile(path) as zf:
            namelist = set(zf.namelist())
            if "word/document.xml" not in namelist:
                raise ValueError(f"DOCX fallback parser could not find word/document.xml in {path.name}") from original_error

            body_root = ET.fromstring(zf.read("word/document.xml"))
            body = body_root.find("w:body", NS)
            if body is not None:
                for child in list(body):
                    local_name = _local_name(child.tag)
                    if local_name == "p":
                        raw_paragraph_count += 1
                        cleaned = _normalize_text(_extract_paragraph_text_xml(child))
                        if cleaned:
                            paragraph_no += 1
                            block_type_counts["body_paragraph"] += 1
                            text_units.append(
                                TextUnit(
                                    document_name=file_info.document_name,
                                    file_path=str(file_info.path),
                                    file_type=file_info.file_type,
                                    locator_type="paragraph",
                                    locator_value=paragraph_no,
                                    text=cleaned,
                                    metadata={"source_zone": "body", "source_kind": "body_paragraph"},
                                )
                            )
                    elif local_name == "tbl":
                        body_table_count += 1
                        cleaned = _normalize_text(_extract_table_text_xml(child))
                        if cleaned:
                            paragraph_no += 1
                            block_type_counts["body_table"] += 1
                            text_units.append(
                                TextUnit(
                                    document_name=file_info.document_name,
                                    file_path=str(file_info.path),
                                    file_type=file_info.file_type,
                                    locator_type="paragraph",
                                    locator_value=paragraph_no,
                                    text=cleaned,
                                    metadata={"source_zone": "body", "source_kind": "body_table"},
                                )
                            )

            header_files = sorted(name for name in namelist if name.startswith("word/header") and name.endswith(".xml"))
            footer_files = sorted(name for name in namelist if name.startswith("word/footer") and name.endswith(".xml"))

            paragraph_no = _append_story_xml_units(
                zf=zf,
                story_files=header_files,
                story_name="header",
                file_info=file_info,
                start_no=paragraph_no,
                text_units=text_units,
                block_type_counts=block_type_counts,
            )
            paragraph_no = _append_story_xml_units(
                zf=zf,
                story_files=footer_files,
                story_name="footer",
                file_info=file_info,
                start_no=paragraph_no,
                text_units=text_units,
                block_type_counts=block_type_counts,
            )

        raw_text = "\n\n".join(unit.text for unit in text_units)
        metadata = self._build_base_metadata(file_info)
        metadata.update(
            {
                "raw_paragraph_count": raw_paragraph_count,
                "non_empty_paragraph_count": len(text_units),
                "body_table_count": body_table_count,
                "section_count": None,
                "block_type_counts": block_type_counts,
                "parse_backend": "zip-xml-fallback",
                "parse_warning": f"python-docx failed and fallback parser was used: {type(original_error).__name__}: {original_error}",
            }
        )

        return ParsedDocument(
            document_name=file_info.document_name,
            file_path=str(file_info.path),
            file_type=file_info.file_type,
            raw_text=raw_text,
            text_units=text_units,
            metadata=metadata,
        )

    def _append_story_blocks(
        self,
        *,
        file_info: FileInfo,
        section: Section,
        story_name: str,
        story_container: _Document | _Cell,
        start_no: int,
        text_units: list[TextUnit],
        block_type_counts: dict[str, int],
        section_index: int,
    ) -> int:
        paragraph_no = start_no
        seen_texts: set[str] = set()
        for block in _iter_block_items(story_container):
            extracted = _extract_block_text(block)
            cleaned = _normalize_text(extracted)
            if not cleaned:
                continue
            story_key = f"{story_name}:{section_index}:{cleaned}"
            if story_key in seen_texts:
                continue
            seen_texts.add(story_key)
            paragraph_no += 1
            block_kind = _block_kind(block, zone=story_name)
            block_type_counts[block_kind] += 1
            text_units.append(
                TextUnit(
                    document_name=file_info.document_name,
                    file_path=str(file_info.path),
                    file_type=file_info.file_type,
                    locator_type="paragraph",
                    locator_value=paragraph_no,
                    text=cleaned,
                    metadata={
                        "source_zone": story_name,
                        "source_kind": block_kind,
                        "section_index": section_index,
                    },
                )
            )
        return paragraph_no


def _iter_block_items(parent) -> Iterator[Paragraph | Table]:
    if isinstance(parent, _Document):
        parent_element = parent.element.body
    elif hasattr(parent, "_tc"):
        parent_element = parent._tc
    elif hasattr(parent, "_element"):
        parent_element = parent._element
    else:
        raise TypeError(f"Unsupported parent type for DOCX block iteration: {type(parent)!r}")
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _extract_block_text(block: Paragraph | Table) -> str:
    if isinstance(block, Paragraph):
        return block.text or ""

    rows: list[str] = []
    for row in block.rows:
        cell_texts = []
        for cell in row.cells:
            pieces = [p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()]
            if pieces:
                cell_texts.append(" ".join(pieces))
        if cell_texts:
            rows.append(" | ".join(cell_texts))
    return "\n".join(rows)


def _block_kind(block: Paragraph | Table, *, zone: str) -> str:
    if isinstance(block, Paragraph):
        return f"{zone}_paragraph"
    return f"{zone}_table"


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines).strip()


def _local_name(tag: str) -> str:
    if "}" in tag:
        return tag.rsplit("}", 1)[1]
    return tag


def _extract_paragraph_text_xml(p_el: ET.Element) -> str:
    parts: list[str] = []
    for node in p_el.iter():
        local = _local_name(node.tag)
        if local == "t":
            if node.text:
                parts.append(node.text)
        elif local in {"tab"}:
            parts.append("\t")
        elif local in {"br", "cr"}:
            parts.append("\n")
    return "".join(parts)


def _extract_table_text_xml(tbl_el: ET.Element) -> str:
    rows: list[str] = []
    for tr in tbl_el.findall(".//w:tr", NS):
        cell_texts: list[str] = []
        for tc in tr.findall("./w:tc", NS):
            pieces: list[str] = []
            for p in tc.findall(".//w:p", NS):
                cleaned = _normalize_text(_extract_paragraph_text_xml(p))
                if cleaned:
                    pieces.append(cleaned)
            if pieces:
                cell_texts.append(" ".join(pieces))
        if cell_texts:
            rows.append(" | ".join(cell_texts))
    return "\n".join(rows)


def _append_story_xml_units(
    *,
    zf: ZipFile,
    story_files: list[str],
    story_name: str,
    file_info: FileInfo,
    start_no: int,
    text_units: list[TextUnit],
    block_type_counts: dict[str, int],
) -> int:
    paragraph_no = start_no
    seen_texts: set[str] = set()
    for story_index, story_file in enumerate(story_files, start=1):
        root = ET.fromstring(zf.read(story_file))
        for child in list(root):
            local_name = _local_name(child.tag)
            if local_name not in {"p", "tbl"}:
                continue
            extracted = _extract_paragraph_text_xml(child) if local_name == "p" else _extract_table_text_xml(child)
            cleaned = _normalize_text(extracted)
            if not cleaned:
                continue
            dedupe_key = f"{story_name}:{story_index}:{cleaned}"
            if dedupe_key in seen_texts:
                continue
            seen_texts.add(dedupe_key)
            paragraph_no += 1
            block_kind = f"{story_name}_{'paragraph' if local_name == 'p' else 'table'}"
            block_type_counts[block_kind] += 1
            text_units.append(
                TextUnit(
                    document_name=file_info.document_name,
                    file_path=str(file_info.path),
                    file_type=file_info.file_type,
                    locator_type="paragraph",
                    locator_value=paragraph_no,
                    text=cleaned,
                    metadata={
                        "source_zone": story_name,
                        "source_kind": block_kind,
                        "story_file": Path(story_file).name,
                    },
                )
            )
    return paragraph_no
